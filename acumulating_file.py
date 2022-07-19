import os
'''
import docx
print(os.getcwd())
os.chdir('D:\\Basic assignment')
post = [i for i in os.listdir()]
print(post)
#name = str(os.path.split(os.getcwd())[1]+' open')
#with  open('Basic_Assignment_1.docx') as f:
from docx import Document
doc = Document('Basic_Assignment_1.docx')
print(doc.paragraphs)
for para in doc.paragraphs:
    print(para.text)


with open('gowith.txt') as outfile:
    file_names = []
    with open('file1.txt') as f:
        for names in filename:
            with open('') as infile:
                outfile.write(infilr.read())
            outfile.write('\n')    
'''
from docx import Document
def filemerge_creater(name,path):            
    try:
        
        os.chdir('D:\workbook')
        print(os.getcwd())
        with open(f'{name}.txt','w+') as f:
            os.chdir(path)
            for i in os.listdir(path):
                print(i)
                docy =Document(str(i))
                f.write(f'\t\t{i} \n')
                for para in docy.paragraphs:
                    f.write(para.text)
                    f.write('\n')   
                f.write('\n\n\n')
    except Exception as e:
        print(e)


filemerge_creater('autoo','D:\Advance_Assignment')    